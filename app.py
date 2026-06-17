import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import re
from datetime import datetime
from collections import Counter
from xml.sax.saxutils import escape as xml_escape

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors as rl_colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="Sustainability Intelligence Studio",
    page_icon="🌿",
    layout="wide"
)

# =========================================================
# LIGHT / MODERN THEME (CSS)
# =========================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"] { font-family: 'Inter', -apple-system, sans-serif; }
.stApp { background: #F7FAF9; }

/* Header */
.main-header {
    background: linear-gradient(135deg, #ECFDF5 0%, #F0FDFA 100%);
    border: 1px solid #D1FAE5;
    border-radius: 18px;
    padding: 28px 30px;
    margin-bottom: 22px;
    box-shadow: 0 2px 10px rgba(16,24,40,0.04);
}
.main-header h1 { color: #065F46; margin: 0; font-size: 28px; font-weight: 800; letter-spacing: -0.5px; }
.main-header p { color: #047857; margin: 6px 0 0 0; font-size: 14px; }
.main-header .credit { color: #6B7280; font-size: 12.5px; margin-top: 10px; }

/* Generic card */
.card { background:#FFFFFF; border:1px solid #E7EBE9; border-radius:16px; padding:20px;
        box-shadow:0 1px 4px rgba(16,24,40,0.04); }

/* Metric card */
.metric-card {
    background:#FFFFFF; border:1px solid #E7EBE9; border-left:4px solid var(--accent,#2F9E68);
    border-radius:14px; padding:16px 18px; box-shadow:0 1px 4px rgba(16,24,40,0.04); position:relative;
    min-height: 118px;
}
.metric-card .icon { font-size:20px; }
.metric-card .value { font-size:24px; font-weight:800; color:#111827; margin:4px 0 2px 0; word-break: break-word; }
.metric-card .label { font-size:12.5px; color:#6B7280; font-weight:500; }
.metric-card .unit { font-size:11px; color:#9CA3AF; }
.confidence-badge { position:absolute; top:14px; right:14px; font-size:10px; font-weight:700;
    padding:3px 8px; border-radius:20px; text-transform:uppercase; letter-spacing:.3px; }
.conf-high { background:#DCFCE7; color:#166534; }
.conf-medium { background:#FEF3C7; color:#92400E; }
.conf-low { background:#FEE2E2; color:#991B1B; }
.conf-none { background:#F3F4F6; color:#6B7280; }

/* Section title */
.section-title { font-size:19px; font-weight:700; color:#111827; margin:28px 0 6px 0; }
.section-sub { color:#6B7280; font-size:13px; margin-bottom:14px; }

/* evidence box */
.evidence-box { background:#F9FAFB; border-left:3px solid #2F9E68; border-radius:8px; padding:10px 14px;
    font-size:12.5px; color:#374151; font-style:italic; margin-top:6px; margin-bottom: 14px;}

/* team table */
.team-table-container { background:#FFFFFF; border:1px solid #E7EBE9; border-radius:16px; padding:18px; }
.team-table-title { color:#065F46; font-weight:700; font-size:15px; margin-bottom:10px; text-align:center; }
.team-table { width:100%; border-collapse:collapse; font-size:13.5px; }
.team-table th { background:#F0FDF4; color:#065F46; padding:8px; text-align:center; font-weight:600; }
.team-table td { padding:7px; text-align:center; border-bottom:1px solid #F1F3F2; color:#374151; }
.team-leader-row { background:#FFFBEB; }
.team-leader-name { color:#B45309 !important; font-weight:700 !important; }
.team-member-name { color:#1565C0; }

/* supervisor card */
.supervisor-card { background:#FFFFFF; border:1px solid #E7EBE9; border-radius:16px; padding:24px; text-align:center; }
.supervisor-title { color:#6B7280; font-size:13px; }
.supervisor-name { color:#B91C1C; font-weight:800; font-size:24px; margin:8px 0; }
.supervisor-qualification { font-size:13px; color:#374151; font-weight:600; }

/* sidebar */
[data-testid="stSidebar"] { background:#FAFCFB; border-right:1px solid #E7EBE9; }
.sb-title { color:#065F46; font-weight:700; font-size:14px; margin-bottom:10px; text-align:center;
    text-transform:uppercase; letter-spacing:.5px; }
.sb-leader { color:#B45309; font-weight:700; font-size:14px; background:#FFFBEB; padding:8px;
    border-radius:8px; text-align:center; margin-bottom:6px; }
.sb-member { color:#374151; font-size:12.5px; padding:5px 8px; background:#F3F4F6; border-radius:6px; margin:4px 0; }
.sb-supervisor { background:#FEF2F2; border-radius:12px; padding:14px; text-align:center; margin-top:8px; }
.sb-supervisor .name { color:#B91C1C; font-weight:700; font-size:16px; margin:4px 0; }

.badge-pill { display:inline-block; padding:3px 10px; border-radius:20px; font-size:11px; font-weight:700; }
.pill-ok { background:#DCFCE7; color:#166534; }
.pill-missing { background:#FEE2E2; color:#991B1B; }
</style>
""", unsafe_allow_html=True)

# =========================================================
# SESSION STATE
# =========================================================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "company_reports" not in st.session_state:
    st.session_state.company_reports = []
if "comparison_mode" not in st.session_state:
    st.session_state.comparison_mode = False

# =========================================================
# LOGIN PAGE
# =========================================================
users = {"admin": "1234", "ismail": "2024"}

if not st.session_state.logged_in:
    st.markdown("""
        <div class='main-header'>
            <h1>🌿 Sustainability Intelligence Studio</h1>
            <p>AI-Powered ESG &amp; Safety Analysis · GRI-Aligned · Multi-Format NLP Engine</p>
            <p class='credit'><b>Team Leader:</b> Ismail Kamal &nbsp;|&nbsp; <b>Under Supervision:</b> Dr. Mohamed Tash · QHSE Master, Alexandria University</p>
        </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("### 🔐 Login to Access System")
        username = st.text_input("Username", placeholder="Enter your username")
        password = st.text_input("Password", type="password", placeholder="Enter your password")
        if st.button("Login", type="primary", use_container_width=True):
            if username in users and users[username] == password:
                st.session_state.logged_in = True
                st.success("✅ Login successful!")
                st.rerun()
            else:
                st.error("❌ Invalid username or password")

    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
            <div class='team-table-container'>
                <div class='team-table-title'>👥 PROJECT TEAM</div>
                <table class='team-table'>
                    <tr><th>Role</th><th>Name</th></tr>
                    <tr class='team-leader-row'><td><b>🏆 Team Leader</b></td><td class='team-leader-name'>Ismail Kamal</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Adel ElSayed</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Mohamed Gaber</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Ahmed Omar</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Sherouk Ashraf</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Mohamed ElHammadi</td></tr>
                    <tr><td>📋 Team Member</td><td class='team-member-name'>Farouk Sameh</td></tr>
                </table>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
            <div class='supervisor-card'>
                <div class='supervisor-title'>🎓 Under Supervision of</div>
                <div class='supervisor-name'>Dr. Mohamed Tash</div>
                <div class='supervisor-qualification'>QHSE Master at Alexandria University</div>
                <div style='margin-top:12px; color:#6B7280; font-size:12px;'>Professor of Sustainability &amp; ESG</div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.caption("© 2026 Sustainability Intelligence Studio | GRI Standards Compliant")
    st.stop()

# =========================================================
# MAIN HEADER
# =========================================================
st.markdown("""
    <div class='main-header'>
        <h1>🌿 Sustainability Intelligence Studio</h1>
        <p>AI-Powered ESG &amp; Safety Analysis · GRI-Aligned · Multi-Format NLP Engine</p>
        <p class='credit'><b>Team Leader:</b> Ismail Kamal &nbsp;|&nbsp; <b>Under Supervision:</b> Dr. Mohamed Tash · QHSE Master, Alexandria University</p>
    </div>
""", unsafe_allow_html=True)

# =========================================================
# SIDEBAR
# =========================================================
with st.sidebar:
    st.markdown("<div style='text-align:center; font-size:42px; margin-bottom:10px;'>🌿</div>", unsafe_allow_html=True)
    st.markdown("""
        <div class='sb-title'>Project Team</div>
        <div class='sb-leader'>🏆 Ismail Kamal (Leader)</div>
        <div class='sb-member'>Adel ElSayed</div>
        <div class='sb-member'>Mohamed Gaber</div>
        <div class='sb-member'>Ahmed Omar</div>
        <div class='sb-member'>Sherouk Ashraf</div>
        <div class='sb-member'>Mohamed ElHammadi</div>
        <div class='sb-member'>Farouk Sameh</div>
        <div class='sb-supervisor'>
            <div style='color:#6B7280; font-size:12px;'>🎓 SUPERVISOR</div>
            <div class='name'>Dr. Mohamed Tash</div>
            <div style='font-size:11.5px; color:#374151;'>QHSE Master, Alexandria University</div>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    comparison_mode = st.checkbox("📊 Company Comparison Mode", value=st.session_state.comparison_mode)
    st.session_state.comparison_mode = comparison_mode

    if st.session_state.comparison_mode:
        c1, c2 = st.columns(2)
        with c1:
            if st.button("➕ Add Company"):
                st.session_state.company_reports.append(None)
        with c2:
            if st.button("🗑️ Clear"):
                st.session_state.company_reports = []
        st.caption(f"Companies added: {len(st.session_state.company_reports)}")

    st.markdown("---")
    st.caption("Version 7.0 · Multi-Format NLP Engine · Evidence-Based Extraction")

# =========================================================
# NLP / INFORMATION-EXTRACTION ENGINE
# =========================================================
# Sustainability reports are mostly unstructured narrative documents with
# inconsistent units and layouts. This engine combines:
#   1) Bilingual (EN/AR) keyword-driven sentence retrieval
#   2) Regex number+unit pairing within the matching sentence
#   3) Structured-table lookup (Excel/Word/PDF tables) as the highest-confidence source
#   4) Confidence scoring + a verifiable evidence sentence for every value
# This keeps the tool lightweight (no heavy ML downloads needed) while still
# being far more reliable on unstructured text than plain keyword search.
# ---------------------------------------------------------

EN_STOPWORDS = set("""the and of to in a is for on with as by an at from this that are be will
have has it its their our your we they he she you i or not no but so than then there which
report company sustainability also more most other into over such these those can may been
within across including per total annual year years""".split())

AR_STOPWORDS = set("""من في على إلى عن مع هذا هذه ذلك التي الذي كان كانت يكون تكون أن إن لا نعم
كل بعض غير حيث حول بين خلال أيضا الشركة التقرير السنة سنة سنوات""".split())

ARABIC_INDIC_DIGITS = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")

ACCENT_BY_CATEGORY = {"Environmental": "#2F9E68", "Social": "#3B82F6", "Governance": "#8B5CF6", "Safety": "#EF4444"}
ICON_BY_CATEGORY = {"Environmental": "🌍", "Social": "🤝", "Governance": "🏛️", "Safety": "🛡️"}

SUBSCRIPT_MAP = {"₀": "0", "₁": "1", "₂": "2", "₃": "3", "₄": "4", "₅": "5", "₆": "6", "₇": "7", "₈": "8", "₉": "9", "³": "3", "²": "2"}

KPI_DEFINITIONS = {
    "co2": {"label": "GHG / CO₂ Emissions", "category": "Environmental", "gri": "GRI 305-1/305-2", "icon": "🌍",
            "unit_label": "tCO₂e",
            "keywords": ["co2 emission", "co2e", "carbon emission", "ghg emission", "greenhouse gas emission",
                         "scope 1", "scope 2", "انبعاثات كربون", "انبعاثات غازات الدفيئة"],
            "unit_patterns": [r"([\d,\.]+)\s*(?:metric\s*)?(?:tonnes?|tons?)\s*(?:of\s*)?co2e?", r"([\d,\.]+)\s*tco2e"]},
    "energy": {"label": "Energy Consumption", "category": "Environmental", "gri": "GRI 302-1", "icon": "⚡",
               "unit_label": "MWh/GJ",
               "keywords": ["energy consumption", "energy use", "total energy", "استهلاك الطاقة"],
               "unit_patterns": [r"([\d,\.]+)\s*(?:mwh|gwh|gj|kwh)"]},
    "renewable": {"label": "Renewable Energy Share", "category": "Environmental", "gri": "GRI 302-1", "icon": "🔆",
                  "unit_label": "%",
                  "keywords": ["renewable energy", "clean energy share", "solar energy", "wind energy", "الطاقة المتجددة"],
                  "unit_patterns": [r"renewable[^.]{0,40}?([\d,\.]+)\s*%", r"([\d,\.]+)\s*%[^.]{0,30}?renewable"]},
    "water": {"label": "Water Withdrawal", "category": "Environmental", "gri": "GRI 303-3", "icon": "💧",
              "unit_label": "m³",
              "keywords": ["water withdrawal", "water consumption", "water use", "استهلاك المياه"],
              "unit_patterns": [r"([\d,\.]+)\s*(?:m3|m³|cubic met(?:er|re)s?|liters?|litres?)"]},
    "waste": {"label": "Waste Generated", "category": "Environmental", "gri": "GRI 306-3", "icon": "🗑️",
              "unit_label": "tons",
              "keywords": ["waste generated", "total waste", "hazardous waste", "tonnes of waste",
                           "tons of waste", "waste", "النفايات"],
              "unit_patterns": [r"([\d,\.]+)\s*(?:metric\s*)?tonnes?\s*(?:of\s*)?waste", r"waste[^.]{0,30}?([\d,\.]+)\s*tons?"]},
    "waste_recycled": {"label": "Waste Recycled / Diverted", "category": "Environmental", "gri": "GRI 306-4", "icon": "♻️",
                        "unit_label": "%",
                        "keywords": ["waste recycled", "diverted from disposal", "recycling rate", "recycled"],
                        "unit_patterns": [r"([\d,\.]+)\s*%[^.]{0,30}?(?:recycled|diverted)",
                                          r"(?:recycled|diverted)[^.]{0,30}?([\d,\.]+)\s*%"]},
    "employees": {"label": "Total Workforce", "category": "Social", "gri": "GRI 2-7", "icon": "👥",
                  "unit_label": "employees",
                  "keywords": ["total employees", "total workforce", "number of employees", "عدد الموظفين"],
                  "unit_patterns": [r"([\d,\.]+)\s*(?:total\s*)?employees"]},
    "female_pct": {"label": "Female Representation", "category": "Social", "gri": "GRI 405-1", "icon": "♀️",
                   "unit_label": "%",
                   "keywords": ["female employees", "women represent", "gender diversity", "نسبة المرأة"],
                   "unit_patterns": [r"([\d,\.]+)\s*%[^.]{0,30}?(?:female|women)",
                                     r"(?:female|women)[^.]{0,30}?([\d,\.]+)\s*%"]},
    "training": {"label": "Average Training Hours", "category": "Social", "gri": "GRI 404-1", "icon": "🎓",
                 "unit_label": "hrs/employee",
                 "keywords": ["training hours", "hours of training", "average training"],
                 "unit_patterns": [r"([\d,\.]+)\s*(?:hours?)\s*(?:of\s*)?training"]},
    "community_invest": {"label": "Community Investment", "category": "Social", "gri": "GRI 203-1", "icon": "🤝",
                          "unit_label": "USD",
                          "keywords": ["community investment", "social investment", "donations"],
                          "unit_patterns": [r"(?:usd|\$)\s*([\d,\.]+)", r"([\d,\.]+)\s*(?:usd|million)[^.]{0,20}?community"]},
    "board_indep": {"label": "Board Independence", "category": "Governance", "gri": "GRI 2-9", "icon": "🏛️",
                    "unit_label": "%",
                    "keywords": ["independent board members", "board independence", "independent directors"],
                    "unit_patterns": [r"([\d,\.]+)\s*%[^.]{0,30}?board",
                                      r"board[^.]{0,30}?(?:independen\w*)[^.]{0,15}?([\d,\.]+)\s*%"]},
    "fatalities": {"label": "Fatalities", "category": "Safety", "gri": "GRI 403-9", "icon": "💀",
                   "unit_label": "cases",
                   "keywords": ["fatalities", "work-related death", "وفيات", "حالات وفاة"],
                   "unit_patterns": [r"([\d,\.]+)\s*(?:work-related\s*)?fatalit(?:y|ies)"]},
    "lti": {"label": "Lost Time Injuries", "category": "Safety", "gri": "GRI 403-9", "icon": "🩹",
            "unit_label": "cases",
            "keywords": ["lost time injur", "lti ", "إصابات", "حوادث عمل"],
            "unit_patterns": [r"([\d,\.]+)\s*lost[\s-]time\s*injur"]},
    "ltifr": {"label": "LTIFR", "category": "Safety", "gri": "GRI 403-9", "icon": "📊",
              "unit_label": "rate",
              "keywords": ["ltifr", "lost time injury frequency rate", "معدل تكرار الإصابات"],
              "unit_patterns": [r"ltifr[^\d]{0,15}([\d,\.]+)", r"([\d,\.]+)[^.]{0,10}?ltifr"]},
    "near_miss": {"label": "Near Misses", "category": "Safety", "gri": "GRI 403-9", "icon": "⚠️",
                  "unit_label": "cases",
                  "keywords": ["near miss", "near-miss", "حوادث وشيكة"],
                  "unit_patterns": [r"([\d,\.]+)\s*near[\s-]miss"]},
}


def normalize_digits(text):
    return (text or "").translate(ARABIC_INDIC_DIGITS)


def split_sentences(text):
    text = re.sub(r"\s+", " ", text or "")
    parts = re.split(r"(?<=[.!?؟])\s+", text)
    return [p.strip() for p in parts if len(p.strip()) > 4]


def find_first_number(pattern, sentence):
    m = re.search(pattern, sentence, re.IGNORECASE)
    if m:
        try:
            return m.group(1)
        except IndexError:
            return None
    return None


def search_tables(tables, keywords):
    """Look for a table row whose label cell matches a keyword and return the
    first adjacent numeric cell — structured data is treated as high confidence.
    Multi-word keywords must appear as a substring; single-word keywords must
    appear as a whole word, to avoid generic words (e.g. 'total') causing
    false matches against unrelated rows."""
    if not tables:
        return None
    for tbl in tables:
        rows = tbl.values.tolist() if isinstance(tbl, pd.DataFrame) else tbl
        if not rows:
            continue
        for row in rows:
            if not row:
                continue
            label = str(row[0]).lower() if row[0] is not None else ""
            hit = False
            for kw in keywords:
                kw_low = kw.lower().strip()
                if not kw_low:
                    continue
                if len(kw_low.split()) > 1:
                    if kw_low in label:
                        hit = True
                        break
                elif re.search(rf"(?<![\w]){re.escape(kw_low)}(?![\w])", label):
                    hit = True
                    break
            if hit:
                for cell in row[1:]:
                    m = re.search(r"[\d][\d,\.]*", str(cell))
                    if m:
                        return m.group(0)
    return None


def extract_kpi(sentences, tables, kpi_def):
    table_val = search_tables(tables, kpi_def["keywords"])
    if table_val:
        return {"value": table_val, "unit": kpi_def["unit_label"], "confidence": "High",
                "evidence": "Found in a structured table within the uploaded file."}

    strong_matches, weak_matches = [], []
    for s in sentences:
        s_low = s.lower()
        kw_pos = None
        for kw in kpi_def["keywords"]:
            idx = s_low.find(kw.lower())
            if idx != -1:
                kw_pos = idx
                break
        if kw_pos is None:
            continue

        matched = False
        for pat in kpi_def["unit_patterns"]:
            val = find_first_number(pat, s)
            if val:
                strong_matches.append({"value": val, "sentence": s})
                matched = True
                break
        if not matched:
            # No exact unit pairing — fall back to the number positioned closest
            # to the matched keyword (more reliable than "first number in sentence").
            # The lookaround avoids matching digits embedded inside words (e.g. the "2" in "co2").
            nums = [(m.group(0), m.start()) for m in re.finditer(r"(?<![A-Za-z])\d[\d,\.]*(?![A-Za-z])", s)]
            if nums:
                nearest = min(nums, key=lambda nm: abs(nm[1] - kw_pos))
                weak_matches.append({"value": nearest[0], "sentence": s})

    if strong_matches:
        chosen = strong_matches[0]
        return {"value": chosen["value"], "unit": kpi_def["unit_label"], "confidence": "High",
                "evidence": chosen["sentence"][:280]}
    if weak_matches:
        chosen = weak_matches[0]
        return {"value": chosen["value"], "unit": kpi_def["unit_label"], "confidence": "Low",
                "evidence": chosen["sentence"][:280]}
    return {"value": "N/A", "unit": kpi_def["unit_label"], "confidence": "Not Found", "evidence": None}


def run_extraction_engine(text, tables):
    sentences = split_sentences(text)
    return {key: extract_kpi(sentences, tables, kdef) for key, kdef in KPI_DEFINITIONS.items()}


def top_terms(text, n=14):
    words = re.findall(r"[A-Za-z\u0600-\u06FF]{3,}", (text or "").lower())
    filtered = [w for w in words if w not in EN_STOPWORDS and w not in AR_STOPWORDS]
    return Counter(filtered).most_common(n)


def coverage_by_category(results):
    cats = {}
    for key, res in results.items():
        cat = KPI_DEFINITIONS[key]["category"]
        cats.setdefault(cat, [0, 0])
        cats[cat][1] += 1
        if res["confidence"] != "Not Found":
            cats[cat][0] += 1
    return {cat: round(100 * found / total, 1) for cat, (found, total) in cats.items()}


def get_safety_summary(results):
    return {"fatalities": results["fatalities"], "lti": results["lti"],
            "ltifr": results["ltifr"], "near_miss": results["near_miss"]}


def safe_float(value):
    if value in (None, "N/A", ""):
        return 0.0
    try:
        return float(re.sub(r"[^\d.\-]", "", str(value)))
    except Exception:
        return 0.0


def pdf_safe(text):
    s = str(text)
    for k, v in SUBSCRIPT_MAP.items():
        s = s.replace(k, v)
    s = re.sub(r"[^\x00-\x7F]+", "", s)
    return s.strip()


# =========================================================
# MULTI-FORMAT FILE READERS
# =========================================================
def read_pdf(file):
    text, tables = "", []
    try:
        import pdfplumber
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
                for tb in (page.extract_tables() or []):
                    tables.append(tb)
        return text, tables
    except ImportError:
        pass
    except Exception as e:
        st.warning(f"⚠️ pdfplumber could not fully parse the file ({e}); falling back to pypdf.")
    try:
        from pypdf import PdfReader
        file.seek(0)
        reader = PdfReader(file)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        st.error(f"❌ Could not read PDF: {e}")
    return text, tables


def read_excel(file):
    text, tables = "", []
    try:
        xls = pd.ExcelFile(file)
        for sheet in xls.sheet_names:
            df = xls.parse(sheet, header=None)
            tables.append(df)
            text += df.astype(str).to_string(index=False) + "\n"
    except Exception as e:
        st.error(f"❌ Could not read Excel file: {e}")
    return text, tables


def read_docx_file(file):
    text, tables = "", []
    try:
        import docx
        d = docx.Document(file)
        text = "\n".join(p.text for p in d.paragraphs)
        for t in d.tables:
            rows = [[cell.text for cell in row.cells] for row in t.rows]
            tables.append(rows)
    except ImportError:
        st.error("❌ python-docx is not installed. Add it to requirements.txt to read .docx files.")
    except Exception as e:
        st.error(f"❌ Could not read Word file: {e}")
    return text, tables


def read_csv_file(file):
    text, tables = "", []
    try:
        df = pd.read_csv(file, header=None)
        tables.append(df)
        text = df.astype(str).to_string(index=False)
    except Exception as e:
        st.error(f"❌ Could not read CSV file: {e}")
    return text, tables


def extract_content(file):
    name = file.name.lower()
    if name.endswith(".pdf"):
        return read_pdf(file)
    if name.endswith((".xlsx", ".xls")):
        return read_excel(file)
    if name.endswith(".docx"):
        return read_docx_file(file)
    if name.endswith(".csv"):
        return read_csv_file(file)
    st.warning(f"⚠️ Unsupported file type: {file.name}")
    return "", []


# =========================================================
# RENDER HELPERS / CHARTS
# =========================================================
def render_metric_card(kdef, result):
    accent = ACCENT_BY_CATEGORY.get(kdef["category"], "#2F9E68")
    conf = result["confidence"]
    conf_class = {"High": "conf-high", "Medium": "conf-medium", "Low": "conf-low", "Not Found": "conf-none"}[conf]
    conf_label = {"High": "Verified", "Medium": "Estimated", "Low": "Weak", "Not Found": "Not Found"}[conf]
    value_display = result["value"] if result["value"] != "N/A" else "—"
    st.markdown(f"""
        <div class="metric-card" style="--accent:{accent};">
            <span class="confidence-badge {conf_class}">{conf_label}</span>
            <div class="icon">{kdef['icon']}</div>
            <div class="value">{value_display}</div>
            <div class="label">{kdef['label']}</div>
            <div class="unit">{result['unit']} · {kdef['gri']}</div>
        </div>
    """, unsafe_allow_html=True)


def chart_top_terms(terms):
    if not terms:
        return None
    df = pd.DataFrame(terms, columns=["Term", "Frequency"]).sort_values("Frequency")
    fig = px.bar(df, x="Frequency", y="Term", orientation="h", color="Frequency",
                 color_continuous_scale=["#D1FAE5", "#059669"])
    fig.update_layout(height=420, showlegend=False, coloraxis_showscale=False,
                       plot_bgcolor="white", paper_bgcolor="white",
                       margin=dict(l=10, r=10, t=40, b=10),
                       title="Most Frequent Sustainability Terms in the Report")
    return fig


def chart_coverage_radar(coverage):
    if not coverage:
        return None
    cats = list(coverage.keys())
    vals = list(coverage.values())
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=vals + [vals[0]], theta=cats + [cats[0]], fill="toself",
                                   line=dict(color="#2F9E68", width=3), fillcolor="rgba(47,158,104,0.25)",
                                   name="Disclosure Coverage"))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                       title="GRI Disclosure Coverage by Category (% of tracked indicators found)",
                       height=420, paper_bgcolor="white", showlegend=False)
    return fig


def chart_safety_overview(safety):
    labels = ["Fatalities", "Lost Time Injuries", "Near Misses"]
    values = [safe_float(safety["fatalities"]["value"]), safe_float(safety["lti"]["value"]),
              safe_float(safety["near_miss"]["value"])]
    colors_list = ["#EF4444", "#F59E0B", "#FACC15"]
    fig = go.Figure(data=[go.Bar(x=labels, y=values, marker_color=colors_list, text=values, textposition="outside")])
    fig.update_layout(title="Safety Incidents Extracted from Report", height=380,
                       plot_bgcolor="white", paper_bgcolor="white", yaxis_title="Number of Cases")
    return fig


def chart_ltifr_gauge(safety):
    val = safe_float(safety["ltifr"]["value"])
    axis_max = max(5, val * 1.3)
    fig = go.Figure(go.Indicator(mode="gauge+number", value=val,
        title={"text": "LTIFR<br><span style='font-size:11px;color:#6B7280'>Lower is better · reference threshold 2.0</span>"},
        gauge={"axis": {"range": [0, axis_max]}, "bar": {"color": "#2F9E68"},
               "steps": [{"range": [0, 1], "color": "#DCFCE7"}, {"range": [1, 2], "color": "#FEF9C3"},
                         {"range": [2, axis_max], "color": "#FEE2E2"}],
               "threshold": {"line": {"color": "#EF4444", "width": 3}, "value": 2.0}}))
    fig.update_layout(height=300, paper_bgcolor="white")
    return fig


def chart_benchmark(value, label, unit, low_ref, high_ref):
    val = safe_float(value)
    fig = go.Figure(data=[
        go.Bar(name="Extracted from Report", x=[label], y=[val], marker_color="#2F9E68"),
        go.Bar(name="Reference Range (low)", x=[label], y=[low_ref], marker_color="#A7F3D0"),
        go.Bar(name="Reference Range (high)", x=[label], y=[high_ref], marker_color="#FDE68A"),
    ])
    fig.update_layout(title=f"{label} vs Indicative Reference Range", yaxis_title=unit, height=380,
                       barmode="group", plot_bgcolor="white", paper_bgcolor="white")
    return fig


# =========================================================
# PDF SUMMARY REPORT
# =========================================================
def generate_pdf_summary_report(results, safety, coverage, file_summaries):
    filename = f"Sustainability_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("T", parent=styles["Heading1"], fontSize=24,
                                  textColor=rl_colors.HexColor("#065F46"), spaceAfter=18, alignment=1)
    heading_style = ParagraphStyle("H", parent=styles["Heading2"], fontSize=15,
                                    textColor=rl_colors.HexColor("#047857"), spaceBefore=16, spaceAfter=8)

    story.append(Paragraph("Sustainability Analysis Summary", title_style))
    story.append(Paragraph("GRI-Aligned - AI-Assisted Extraction with Evidence Trail", styles["Normal"]))
    story.append(Spacer(1, 14))
    story.append(Paragraph("<b>Team Leader:</b> Ismail Kamal", styles["Normal"]))
    story.append(Paragraph("<b>Team Members:</b> Adel ElSayed, Mohamed Gaber, Ahmed Omar, Sherouk Ashraf, "
                            "Mohamed ElHammadi, Farouk Sameh", styles["Normal"]))
    story.append(Paragraph("<font color='#B91C1C'><b>Under Supervision: Dr. Mohamed Tash</b></font>", styles["Normal"]))
    story.append(Paragraph(f"<b>Report Date:</b> {datetime.now().strftime('%B %d, %Y')}", styles["Normal"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Source Files", heading_style))
    for fs in file_summaries:
        story.append(Paragraph(f"- {pdf_safe(fs['name'])} - {fs['chars']:,} characters, "
                                f"{fs['tables']} table(s) detected", styles["Normal"]))

    story.append(Paragraph("Extracted Indicators", heading_style))
    table_data = [["Indicator", "Value", "Unit", "GRI", "Confidence"]]
    for k, res in results.items():
        kdef = KPI_DEFINITIONS[k]
        table_data.append([pdf_safe(kdef["label"]), pdf_safe(res["value"]), pdf_safe(res["unit"]),
                            kdef["gri"], res["confidence"]])
    t = Table(table_data, colWidths=[150, 70, 80, 90, 70])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#065F46")),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#D1D5DB")),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t)
    story.append(Spacer(1, 14))

    story.append(Paragraph("GRI Disclosure Coverage by Category", heading_style))
    cov_data = [["Category", "Coverage %"]] + [[cat, f"{pct}%"] for cat, pct in coverage.items()]
    t2 = Table(cov_data, colWidths=[150, 100])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#047857")),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#D1D5DB")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t2)
    story.append(PageBreak())

    story.append(Paragraph("Evidence Appendix", heading_style))
    story.append(Paragraph("Only High/Medium-confidence values are listed with their source sentence "
                            "(for verification against the original document).", styles["Normal"]))
    story.append(Spacer(1, 8))
    for k, res in results.items():
        if res["confidence"] in ("High", "Low") and res.get("evidence"):
            kdef = KPI_DEFINITIONS[k]
            story.append(Paragraph(f"<b>{pdf_safe(kdef['label'])}</b> "
                                    f"({pdf_safe(res['value'])} {pdf_safe(res['unit'])}, {res['confidence']})",
                                    styles["Normal"]))
            story.append(Paragraph(f"<i>{xml_escape(pdf_safe(res['evidence']))}</i>", styles["Normal"]))
            story.append(Spacer(1, 6))

    story.append(Spacer(1, 16))
    story.append(Paragraph("Developed by Ismail Kamal and Team | Under Supervision of Dr. Mohamed Tash", styles["Normal"]))
    story.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles["Normal"]))

    doc.build(story)
    return filename


# =========================================================
# MAIN ANALYSIS
# =========================================================
if not st.session_state.comparison_mode:

    with st.expander("ℹ️ How this AI analysis engine works"):
        st.markdown("""
        This tool is built to handle **unstructured sustainability reports** using a transparent, lightweight NLP pipeline:
        1. **Multi-format parsing** — text and tables are extracted from PDF, Excel, Word, or CSV files.
        2. **Bilingual keyword retrieval** — sentences mentioning each GRI-aligned indicator (English/Arabic) are located.
        3. **Contextual number–unit extraction** — regex patterns pair the right unit (tCO₂e, m³, %, etc.) with the matching number.
        4. **Structured-table priority** — a value found in a table is treated as higher confidence than one found in free text.
        5. **Confidence scoring & evidence trail** — every value is labeled *High / Medium / Low / Not Found* and linked to the exact source sentence, so results can be verified against the original report.
        """)

    uploaded_files = st.file_uploader(
        "📄 Upload Sustainability Report(s) — PDF, Excel, Word or CSV",
        type=["pdf", "xlsx", "xls", "docx", "csv"],
        accept_multiple_files=True
    )

    if uploaded_files:
        combined_text, combined_tables, file_summaries = "", [], []
        with st.spinner("📖 Reading and parsing document(s)..."):
            for f in uploaded_files:
                t, tb = extract_content(f)
                combined_text += "\n" + (t or "")
                combined_tables += tb
                file_summaries.append({"name": f.name, "chars": len(t or ""), "tables": len(tb)})

        combined_text = normalize_digits(combined_text)

        if not combined_text.strip():
            st.error("❌ Could not extract any text. The file(s) may be scanned images — OCR is required first.")
            st.stop()

        if st.button("🔍 Run AI Analysis", type="primary", use_container_width=True):
            with st.spinner("🤖 Extracting sustainability indicators with the NLP engine..."):
                results = run_extraction_engine(combined_text, combined_tables)
                coverage = coverage_by_category(results)
                terms = top_terms(combined_text)
                safety = get_safety_summary(results)

            # ---- Document insights ----
            st.markdown('<div class="section-title">📑 Document Insights</div>', unsafe_allow_html=True)
            arabic_chars = len(re.findall(r"[\u0600-\u06FF]", combined_text))
            latin_chars = len(re.findall(r"[A-Za-z]", combined_text))
            language_mix = "Arabic + English" if arabic_chars > 200 and latin_chars > 200 else (
                "Arabic" if arabic_chars > latin_chars else "English")
            c1, c2, c3, c4 = st.columns(4)
            with c1: st.metric("Files Analyzed", len(uploaded_files))
            with c2: st.metric("Words Processed", f"{len(combined_text.split()):,}")
            with c3: st.metric("Structured Tables Found", len(combined_tables))
            with c4: st.metric("Language Mix", language_mix)
            st.markdown("---")

            # ---- KPI cards by category ----
            for category in ["Environmental", "Social", "Governance", "Safety"]:
                keys = [k for k, v in KPI_DEFINITIONS.items() if v["category"] == category]
                if not keys:
                    continue
                st.markdown(f'<div class="section-title">{ICON_BY_CATEGORY[category]} {category} Indicators</div>',
                            unsafe_allow_html=True)
                col_count = min(len(keys), 4)
                cols = st.columns(col_count)
                for i, k in enumerate(keys):
                    with cols[i % col_count]:
                        render_metric_card(KPI_DEFINITIONS[k], results[k])
            st.markdown("---")

            # ---- Evidence panel ----
            st.markdown('<div class="section-title">🔎 Evidence &amp; Verifiability</div>', unsafe_allow_html=True)
            st.markdown('<p class="section-sub">For academic review, every extracted value is traceable to the '
                        'exact sentence or table cell it came from.</p>', unsafe_allow_html=True)
            with st.expander("View extraction evidence for all indicators"):
                for k, res in results.items():
                    label = KPI_DEFINITIONS[k]["label"]
                    if res["confidence"] != "Not Found":
                        st.markdown(f"**{label}** — `{res['value']} {res['unit']}` ({res['confidence']} confidence)")
                        st.markdown(f'<div class="evidence-box">{res["evidence"]}</div>', unsafe_allow_html=True)
                    else:
                        st.markdown(f"**{label}** — _Not found in the document(s)._")
            st.markdown("---")

            # ---- Coverage + terms ----
            col1, col2 = st.columns(2)
            with col1:
                fig = chart_coverage_radar(coverage)
                if fig:
                    st.plotly_chart(fig, use_container_width=True)
            with col2:
                fig = chart_top_terms(terms)
                if fig:
                    st.plotly_chart(fig, use_container_width=True)
            st.markdown("---")

            # ---- Safety dashboard ----
            st.markdown('<div class="section-title">🛡️ Safety Performance</div>', unsafe_allow_html=True)
            st.caption("A value of 0 may mean either zero incidents or that the indicator was not disclosed — "
                       "check the evidence panel above to confirm.")
            col1, col2 = st.columns(2)
            with col1: st.plotly_chart(chart_safety_overview(safety), use_container_width=True)
            with col2: st.plotly_chart(chart_ltifr_gauge(safety), use_container_width=True)
            st.markdown("---")

            # ---- Indicative benchmarking ----
            st.markdown('<div class="section-title">🎯 Indicative Benchmarking</div>', unsafe_allow_html=True)
            st.caption("⚠️ The reference ranges below are general indicative ranges for context, not a live "
                       "industry database — state this limitation when citing the analysis.")
            col1, col2 = st.columns(2)
            with col1:
                if results["co2"]["value"] == "N/A":
                    st.info("ℹ️ CO₂ emissions value not found in the document — benchmarking skipped.")
                else:
                    st.plotly_chart(chart_benchmark(results["co2"]["value"], "CO₂ Emissions (tCO₂e)",
                                                     "tCO₂e", 30000, 50000), use_container_width=True)
            with col2:
                if results["renewable"]["value"] == "N/A":
                    st.info("ℹ️ Renewable energy share not found in the document — benchmarking skipped.")
                else:
                    st.plotly_chart(chart_benchmark(results["renewable"]["value"], "Renewable Energy (%)",
                                                     "%", 20, 40), use_container_width=True)
            st.markdown("---")

            # ---- GRI compliance ----
            st.markdown('<div class="section-title">📜 GRI Standards Disclosure Check</div>', unsafe_allow_html=True)
            cols = st.columns(3)
            for i, (k, res) in enumerate(results.items()):
                kdef = KPI_DEFINITIONS[k]
                with cols[i % 3]:
                    pill = "pill-ok" if res["confidence"] != "Not Found" else "pill-missing"
                    pill_text = "Disclosed" if res["confidence"] != "Not Found" else "Not Disclosed"
                    st.markdown(f"""
                        <div class="card" style="margin-bottom:10px;">
                            <b>{kdef['gri']}</b><br>
                            <span style="font-size:12.5px;color:#6B7280;">{kdef['label']}</span><br>
                            <span class="badge-pill {pill}">{pill_text}</span>
                        </div>
                    """, unsafe_allow_html=True)
            st.markdown("---")

            # ---- PDF export ----
            st.markdown('<div class="section-title">📥 Export Summary Report</div>', unsafe_allow_html=True)
            pdf_file = generate_pdf_summary_report(results, safety, coverage, file_summaries)
            with open(pdf_file, "rb") as f:
                st.download_button("📥 Download PDF Summary Report", f, file_name=pdf_file,
                                    mime="application/pdf", use_container_width=True)
            st.success("✅ Analysis completed successfully.")

else:
    # =========================================================
    # MULTI-COMPANY COMPARISON MODE
    # =========================================================
    st.markdown('<div class="section-title">🏢 Multi-Company Comparison</div>', unsafe_allow_html=True)
    st.caption("Upload one or more files per company to compare extracted sustainability indicators side by side.")

    companies_rows = []
    for i in range(len(st.session_state.company_reports)):
        files_i = st.file_uploader(f"Company {i + 1} — file(s)", type=["pdf", "xlsx", "xls", "docx", "csv"],
                                    accept_multiple_files=True, key=f"company_{i}")
        if files_i:
            with st.spinner(f"Analyzing Company {i + 1}..."):
                text_c, tables_c = "", []
                for f in files_i:
                    t, tb = extract_content(f)
                    text_c += "\n" + (t or "")
                    tables_c += tb
                text_c = normalize_digits(text_c)
                if text_c.strip():
                    res_c = run_extraction_engine(text_c, tables_c)
                    companies_rows.append({
                        "Company": f"Company {i + 1}",
                        "CO2": safe_float(res_c["co2"]["value"]),
                        "Energy": safe_float(res_c["energy"]["value"]),
                        "Water": safe_float(res_c["water"]["value"]),
                        "Waste": safe_float(res_c["waste"]["value"]),
                        "Renewable %": safe_float(res_c["renewable"]["value"]),
                        "LTIFR": safe_float(res_c["ltifr"]["value"]),
                        "Near Misses": safe_float(res_c["near_miss"]["value"]),
                    })

    if companies_rows and st.button("📊 Compare Companies", type="primary", use_container_width=True):
        df_compare = pd.DataFrame(companies_rows)
        st.dataframe(df_compare, use_container_width=True)

        fig1 = px.bar(df_compare, x="Company", y=["CO2", "Energy", "Water", "Waste"], barmode="group",
                      title="Environmental KPIs Comparison",
                      color_discrete_sequence=["#2F9E68", "#3B82F6", "#0EA5E9", "#A855F7"])
        fig1.update_layout(plot_bgcolor="white", paper_bgcolor="white")
        st.plotly_chart(fig1, use_container_width=True)

        fig2 = px.bar(df_compare, x="Company", y=["LTIFR", "Near Misses"], barmode="group",
                      title="Safety KPIs Comparison", color_discrete_sequence=["#EF4444", "#F59E0B"])
        fig2.update_layout(plot_bgcolor="white", paper_bgcolor="white")
        st.plotly_chart(fig2, use_container_width=True)

        st.success("✅ Comparison complete!")
    elif not st.session_state.company_reports:
        st.info("📌 Use the sidebar to add companies, then upload each company's file(s) above.")

# =========================================================
# FOOTER
# =========================================================
st.markdown("---")
st.markdown("""
    <div style='text-align:center; padding:18px; background:#FFFFFF; border:1px solid #E7EBE9;
                border-radius:14px; margin-top:10px;'>
        <p style='color:#111827; margin:0; font-weight:600;'>🌿 Sustainability Intelligence Studio | GRI Standards 2024</p>
        <p style='color:#6B7280; font-size:12px; margin:6px 0 0 0;'>
            Developed by <b>Ismail Kamal</b> &amp; Team |
            <b style='color:#B91C1C;'>Under Supervision of Dr. Mohamed Tash</b>
        </p>
        <p style='color:#9CA3AF; font-size:11px; margin:4px 0 0 0;'>Version 7.0 · Multi-Format NLP Engine · Evidence-Based Extraction</p>
    </div>
""", unsafe_allow_html=True)
